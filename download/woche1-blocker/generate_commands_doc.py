#!/usr/bin/env python3
"""
Generate B) Commands + Migrationsskripte DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, code, language=""):
    """Add a formatted code block"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'DejaVu Sans Mono'
    run.font.size = Pt(9)
    # Add shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    para._p.get_or_add_pPr().append(shading)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

def create_commands_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('Woche 1 - Commands & Migrationsskripte', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('CargoBit Payment System - Produktionsreife Implementierung')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Table of Contents style header
    doc.add_heading('Inhalt', level=1)
    toc_items = [
        '1. PostgreSQL Migration',
        '2. Redis Rate Limiting',
        '3. Backups + PITR',
        '4. Stripe Webhook Validation',
        '5. Audit Log Hardening',
        '6. Environment Variables',
        '7. Testing Commands'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 1. PostgreSQL Migration
    doc.add_heading('1. PostgreSQL Migration', level=1)
    
    doc.add_heading('1.1 Postgres Instanz erstellen (Neon)', level=2)
    doc.add_paragraph('Option A: Neon (empfohlen, kostenloses Tier)')
    add_code_block(doc, '''# 1. Account erstellen: https://neon.tech
# 2. Project erstellen: "cargobit-production"
# 3. Connection String kopieren

# Connection String Format:
postgresql://[user]:[password]@[endpoint].neon.tech/[database]?sslmode=require''')
    
    doc.add_paragraph('Option B: Supabase Alternative')
    add_code_block(doc, '''# 1. Account erstellen: https://supabase.com
# 2. Project erstellen
# 3. Database > Connection String > Nodejs

postgresql://postgres.[ref]:[password]@aws-0-[region].pooler.supabase.com:6543/postgres''')
    
    doc.add_heading('1.2 Prisma Migration', level=2)
    add_code_block(doc, '''# .env aktualisieren
DATABASE_URL="postgresql://user:pass@host:5432/cargobit?schema=public"

# Falls bisher SQLite ohne Migrations:
npx prisma migrate diff \\
  --from-url "file:./dev.db" \\
  --to-url "$DATABASE_URL" \\
  --script > migrations/init.sql

# Migration anwenden
psql $DATABASE_URL -f migrations/init.sql

# Prisma Client neu generieren
npx prisma generate

# Seed-Daten (optional)
npx prisma db seed''')
    
    doc.add_heading('1.3 Datenbereinigung vor Migration', level=2)
    add_code_block(doc, '''-- Null-Werte finden und bereinigen
SELECT COUNT(*) FROM transactions WHERE user_id IS NULL;
UPDATE transactions SET user_id = 'unknown' WHERE user_id IS NULL;

-- Duplicate finden
SELECT email, COUNT(*) FROM users GROUP BY email HAVING COUNT(*) > 1;

-- Enum-Type erstellen (Postgres)
CREATE TYPE transaction_status AS ENUM ('pending', 'completed', 'failed');

-- Foreign Key Constraints prüfen
SELECT conname, conrelid::regclass, confrelid::regclass
FROM pg_constraint WHERE contype = 'f';''')
    
    doc.add_heading('1.4 Rollback-Plan', level=2)
    add_code_block(doc, '''# SQLite Backup behalten
cp dev.db dev.db.backup

# Bei Problemen zurück zu SQLite:
# 1. DATABASE_URL auf SQLite setzen
# 2. App neu deployen
# 3. Daten aus Postgres exportieren und in SQLite importieren

# Postgres zu SQLite Export
pg_dump -Fc $DATABASE_URL > postgres_backup.dump

# SQLite Import
sqlite3 dev.db < sqlite_schema.sql
sqlite3 dev.db < sqlite_data.sql''')
    
    doc.add_page_break()
    
    # 2. Redis Rate Limiting
    doc.add_heading('2. Redis Rate Limiting', level=1)
    
    doc.add_heading('2.1 Redis Instanz', level=2)
    add_code_block(doc, '''# Option A: Docker (lokal)
docker run -d --name redis -p 6379:6379 redis:7-alpine

# Option B: Upstash (Cloud, empfohlen)
# 1. Account: https://upstash.com
# 2. Redis Database erstellen
# 3. Connection String kopieren

# .env
REDIS_URL="rediss://default:[password]@[endpoint].upstash.io:6379"''')
    
    doc.add_heading('2.2 Token Bucket Implementation', level=2)
    add_code_block(doc, '''// rateLimiter.ts
import { Redis } from "ioredis";

const redis = new Redis(process.env.REDIS_URL!);

export interface RateLimitResult {
  allowed: boolean;
  remaining: number;
  resetAt: number;
}

export async function rateLimit(
  key: string,
  limit: number,
  windowSec: number
): Promise<RateLimitResult> {
  const now = Date.now();
  const windowStart = now - windowSec * 1000;

  const result = await redis
    .multi()
    .zremrangebyscore(key, 0, windowStart)  // Alte Einträge entfernen
    .zadd(key, now, `${now}-${Math.random()}`)  // Neuer Eintrag
    .zcard(key)  // Anzahl zählen
    .pexpire(key, windowSec * 1000)  // TTL setzen
    .exec();

  const count = result?.[2]?.[1] as number;
  const remaining = Math.max(0, limit - count);

  return {
    allowed: count <= limit,
    remaining,
    resetAt: now + windowSec * 1000,
  };
}

// Middleware
export function rateLimitMiddleware(limit: number, windowSec: number) {
  return async (req: any, res: any, next: any) => {
    const key = `ratelimit:${req.ip}:${req.path}`;
    const result = await rateLimit(key, limit, windowSec);

    res.setHeader("X-RateLimit-Limit", limit);
    res.setHeader("X-RateLimit-Remaining", result.remaining);
    res.setHeader("X-RateLimit-Reset", result.resetAt);

    if (!result.allowed) {
      return res.status(429).json({
        error: "Too Many Requests",
        retryAfter: windowSec,
      });
    }

    next();
  };
}''')
    
    doc.add_heading('2.3 Rate Limits konfigurieren', level=2)
    add_code_block(doc, '''// rateLimitConfig.ts
export const RATE_LIMITS = {
  // Öffentliche API
  "GET /api/public": { limit: 100, window: 60 },      // 100/min
  "POST /api/public": { limit: 20, window: 60 },      // 20/min

  // Auth-Endpoints (strenger)
  "POST /api/auth/login": { limit: 5, window: 300 },  // 5/5min
  "POST /api/auth/register": { limit: 3, window: 3600 }, // 3/h

  // Admin-Endpoints (sehr streng)
  "DELETE /api/admin/*": { limit: 10, window: 60 },   // 10/min

  // Webhooks (nur von Stripe)
  "POST /api/webhooks/stripe": { limit: 1000, window: 60 }, // 1000/min
};''')
    
    doc.add_page_break()
    
    # 3. Backups + PITR
    doc.add_heading('3. Backups + PITR', level=1)
    
    doc.add_heading('3.1 Neon PITR (automatisch)', level=2)
    add_code_block(doc, '''# Neon hat eingebautes PITR:
# - Point-in-Time Recovery bis zu 7 Tage (Free)
# - Bis zu 30 Tage (Pro)

# Wiederherstellung über UI oder CLI:
# 1. Neon Console > Branches
# 2. "Time Travel" zu gewünschtem Zeitpunkt
# 3. Neuen Branch erstellen
# 4. Connection String aktualisieren''')
    
    doc.add_heading('3.2 Manual Backup Script', level=2)
    add_code_block(doc, '''#!/bin/bash
# backup.sh - Täglich um 3:00 UTC via Cron

set -e

DATE=$(date +%F)
BACKUP_DIR="/backups"
S3_BUCKET="s3://cargobit-backups/postgres"

# 1. PostgreSQL Backup (compressed)
pg_dump -Fc $DATABASE_URL > "$BACKUP_DIR/cargobit_$DATE.dump"

# 2. Zusätzlich SQL (lesbar)
pg_dump $DATABASE_URL > "$BACKUP_DIR/cargobit_$DATE.sql"

# 3. Zu S3 hochladen (verschlüsselt)
aws s3 cp "$BACKUP_DIR/cargobit_$DATE.dump" \\
  "$S3_BUCKET/daily/" \\
  --storage-class STANDARD_IA \\
  --server-side-encryption aws:kms

# 4. Lokale Backups älter als 7 Tage löschen
find $BACKUP_DIR -name "*.dump" -mtime +7 -delete
find $BACKUP_DIR -name "*.sql" -mtime +7 -delete

echo "Backup completed: $DATE"''')
    
    doc.add_heading('3.3 Restore Script', level=2)
    add_code_block(doc, '''#!/bin/bash
# restore.sh - Restore from backup

set -e

BACKUP_FILE=${1:-"latest"}

if [ "$BACKUP_FILE" = "latest" ]; then
  # Neuestes Backup von S3 holen
  BACKUP_FILE=$(aws s3 ls s3://cargobit-backups/postgres/daily/ | \\
    sort | tail -1 | awk '{print $4}')
  aws s3 cp "s3://cargobit-backups/postgres/daily/$BACKUP_FILE" /tmp/restore.dump
  BACKUP_FILE="/tmp/restore.dump"
fi

# WARNUNG: Dies überschreibt die bestehende Datenbank!
read -p "Are you sure? This will DELETE existing data! (yes/no): " confirm

if [ "$confirm" = "yes" ]; then
  # 1. Alle Verbindungen trennen
  psql $DATABASE_URL -c "SELECT pg_terminate_backend(pid) FROM pg_stat_activity WHERE datname = 'cargobit';"

  # 2. Datenbank löschen und neu erstellen
  psql $DATABASE_URL -c "DROP DATABASE IF EXISTS cargobit;"
  psql $DATABASE_URL -c "CREATE DATABASE cargobit;"

  # 3. Restore
  pg_restore -d $DATABASE_URL $BACKUP_FILE

  echo "Restore completed!"
else
  echo "Aborted."
fi''')
    
    doc.add_page_break()
    
    # 4. Stripe Webhook Validation
    doc.add_heading('4. Stripe Webhook Validation', level=1)
    
    doc.add_heading('4.1 Webhook Endpoint', level=2)
    add_code_block(doc, '''// webhookHandler.ts
import Stripe from "stripe";
import { db } from "./db";

const stripe = new Stripe(process.env.STRIPE_SECRET_KEY!);
const WEBHOOK_SECRET = process.env.STRIPE_WEBHOOK_SECRET!;

// Raw Body ist kritisch!
app.post(
  "/api/webhooks/stripe",
  express.raw({ type: "application/json" }),
  async (req, res) => {
    const sig = req.headers["stripe-signature"];

    let event: Stripe.Event;
    try {
      // Signature Validation - IMMER!
      event = stripe.webhooks.constructEvent(
        req.body, // Raw body, nicht geparst!
        sig as string,
        WEBHOOK_SECRET
      );
    } catch (err: any) {
      console.error("Webhook signature verification failed:", err.message);
      return res.status(400).send(`Webhook Error: ${err.message}`);
    }

    // Idempotency Check
    const processed = await db.processedEvents.findUnique({
      where: { id: event.id },
    });

    if (processed) {
      console.log("Event already processed:", event.id);
      return res.json({ received: true, duplicate: true });
    }

    // Event verarbeiten
    try {
      await handleEvent(event);

      // Als verarbeitet markieren
      await db.processedEvents.create({
        data: {
          id: event.id,
          type: event.type,
          processedAt: new Date(),
        },
      });

      res.json({ received: true });
    } catch (err) {
      console.error("Event processing failed:", err);
      // 5xx = Stripe wiederholt
      // 4xx = Stripe wiederholt nicht
      res.status(500).send("Processing failed");
    }
  }
);

async function handleEvent(event: Stripe.Event) {
  switch (event.type) {
    case "payment_intent.succeeded":
      const payment = event.data.object as Stripe.PaymentIntent;
      await updateOrderStatus(payment.id, "completed");
      break;

    case "payment_intent.payment_failed":
      const failed = event.data.object as Stripe.PaymentIntent;
      await updateOrderStatus(failed.id, "failed");
      break;

    // ... weitere Events
  }
}''')
    
    doc.add_heading('4.2 Webhook Testing', level=2)
    add_code_block(doc, '''# Stripe CLI installieren
brew install stripe/stripe-cli/stripe

# Login
stripe login

# Webhooks lokal testen
stripe listen --forward-to localhost:3000/api/webhooks/stripe

# Test-Event senden
stripe trigger payment_intent.succeeded

# Signature Secret aus der CLI-Ausgabe kopieren
# whsec_xxx -> STRIPE_WEBHOOK_SECRET''')
    
    doc.add_page_break()
    
    # 5. Audit Log Hardening
    doc.add_heading('5. Audit Log Hardening', level=1)
    
    doc.add_heading('5.1 SQL Schema', level=2)
    add_code_block(doc, '''-- audit_log.sql
CREATE TABLE audit_log (
  id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
  timestamp TIMESTAMPTZ NOT NULL DEFAULT now(),
  actor TEXT NOT NULL,
  actor_type TEXT NOT NULL CHECK (actor_type IN ('user', 'system', 'api')),
  action TEXT NOT NULL,
  entity TEXT NOT NULL,
  entity_id TEXT NOT NULL,
  metadata JSONB DEFAULT '{}',
  prev_hash TEXT,
  hash TEXT NOT NULL,
  CONSTRAINT valid_hash CHECK (hash IS NOT NULL)
);

-- Write-Only Constraint via Trigger
CREATE OR REPLACE FUNCTION prevent_audit_modification()
RETURNS TRIGGER AS $$
BEGIN
  RAISE EXCEPTION 'Audit logs are immutable. INSERT only.';
END;
$$ LANGUAGE plpgsql;

CREATE TRIGGER audit_no_update_delete
  BEFORE UPDATE OR DELETE ON audit_log
  FOR EACH ROW EXECUTE FUNCTION prevent_audit_modification();

-- Index für schnelle Abfragen
CREATE INDEX idx_audit_timestamp ON audit_log(timestamp DESC);
CREATE INDEX idx_audit_actor ON audit_log(actor);
CREATE INDEX idx_audit_entity ON audit_log(entity, entity_id);''')
    
    doc.add_heading('5.2 Hash-Chain Implementation', level=2)
    add_code_block(doc, '''// auditLogger.ts
import { db } from "./db";
import crypto from "crypto";

interface AuditEntry {
  actor: string;
  actorType: "user" | "system" | "api";
  action: string;
  entity: string;
  entityId: string;
  metadata?: Record<string, any>;
}

export async function logAudit(entry: AuditEntry): Promise<void> {
  // Letzten Eintrag holen
  const lastEntry = await db.auditLog.findFirst({
    orderBy: { timestamp: "desc" },
    select: { hash: true },
  });

  // Hash berechnen
  const dataToHash = JSON.stringify({
    ...entry,
    timestamp: new Date().toISOString(),
    prevHash: lastEntry?.hash || "genesis",
  });

  const hash = crypto.createHash("sha256").update(dataToHash).digest("hex");

  // Eintragen
  await db.auditLog.create({
    data: {
      ...entry,
      prevHash: lastEntry?.hash || null,
      hash,
    },
  });
}

// Integrität prüfen
export async function verifyAuditIntegrity(): Promise<boolean> {
  const entries = await db.auditLog.findMany({
    orderBy: { timestamp: "asc" },
  });

  let prevHash = "genesis";
  for (const entry of entries) {
    const dataToHash = JSON.stringify({
      actor: entry.actor,
      actorType: entry.actorType,
      action: entry.action,
      entity: entry.entity,
      entityId: entry.entityId,
      metadata: entry.metadata,
      timestamp: entry.timestamp.toISOString(),
      prevHash,
    });

    const expectedHash = crypto
      .createHash("sha256")
      .update(dataToHash)
      .digest("hex");

    if (entry.hash !== expectedHash) {
      console.error("Audit integrity violation at:", entry.id);
      return false;
    }

    prevHash = entry.hash;
  }

  return true;
}''')
    
    doc.add_heading('5.3 S3 Export', level=2)
    add_code_block(doc, '''#!/bin/bash
# export_audit_logs.sh - Täglich um 4:00 UTC

DATE=$(date +%F)
S3_BUCKET="s3://cargobit-audit-logs"

# Audit Logs exportieren (verlüsselt)
psql $DATABASE_URL -c "\\COPY (
  SELECT * FROM audit_log 
  WHERE timestamp >= CURRENT_DATE - INTERVAL '1 day'
  ORDER BY timestamp
) TO STDOUT WITH CSV HEADER" | gzip | \\
aws s3 cp - "$S3_BUCKET/daily/audit_$DATE.csv.gz" \\
  --storage-class GLACIER \\
  --server-side-encryption aws:kms

echo "Audit export completed: $DATE"''')
    
    doc.add_page_break()
    
    # 6. Environment Variables
    doc.add_heading('6. Environment Variables', level=1)
    add_code_block(doc, '''# .env.example - Alle benötigten Variablen

# Database
DATABASE_URL="postgresql://user:pass@host:5432/cargobit?schema=public"
DATABASE_POOL_SIZE="10"

# Redis
REDIS_URL="rediss://default:password@endpoint.upstash.io:6379"

# Stripe
STRIPE_SECRET_KEY="sk_live_xxx"
STRIPE_PUBLISHABLE_KEY="pk_live_xxx"
STRIPE_WEBHOOK_SECRET="whsec_xxx"

# S3/Backup
AWS_ACCESS_KEY_ID="AKIAxxx"
AWS_SECRET_ACCESS_KEY="xxx"
AWS_REGION="eu-central-1"
BACKUP_S3_BUCKET="s3://cargobit-backups"
AUDIT_S3_BUCKET="s3://cargobit-audit-logs"

# Security
JWT_SECRET="your-super-secret-key-min-32-chars"
ENCRYPTION_KEY="32-byte-encryption-key-hex"

# Rate Limiting
RATE_LIMIT_ENABLED="true"
RATE_LIMIT_DEFAULT_LIMIT="100"
RATE_LIMIT_DEFAULT_WINDOW="60"''')
    
    # 7. Testing Commands
    doc.add_heading('7. Testing Commands', level=1)
    
    doc.add_heading('7.1 Integration Tests', level=2)
    add_code_block(doc, '''# Alle Tests ausführen
npm test

# Nur Rate Limiting Tests
npm test -- --grep "rate limit"

# Mit Coverage
npm test -- --coverage

# Watch Mode
npm test -- --watch''')
    
    doc.add_heading('7.2 Load Tests (k6)', level=2)
    add_code_block(doc, '''# k6 installieren
brew install k6

# Load Test ausführen
k6 run load-tests/api-load.js

# Mit 100 gleichzeitigen Usern
k6 run --vus 100 --duration 30s load-tests/api-load.js

# Rate Limit Test
k6 run load-tests/rate-limit.js''')
    
    doc.add_heading('7.3 Security Scan', level=2)
    add_code_block(doc, '''# OWASP ZAP Scan
docker run -t owasp/zap2docker-stable zap-baseline.py \\
  -t https://staging.cargobit.com

# npm audit
npm audit

# Dependencies prüfen
npm outdated

# SQL Injection Test (sqlmap)
sqlmap -u "https://staging.cargobit.com/api/users?id=1" \\
  --batch --level=3''')
    
    # Save
    output_path = '/home/z/my-project/download/woche1-blocker/B_commands_migration_scripts.docx'
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == "__main__":
    create_commands_doc()
